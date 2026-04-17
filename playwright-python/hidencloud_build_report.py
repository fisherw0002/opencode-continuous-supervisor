from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches

ARTIFACTS = Path('/root/.openclaw/workspace/playwright-python/artifacts/hidencloud-renew-2026-04-19')
DOCX_PATH = ARTIFACTS / 'hidencloud-renew-report-2026-04-19.docx'
LOG_PATH = ARTIFACTS / 'run.log'
IMAGES = [
    ('目标页', ARTIFACTS / '01-target-page.png'),
    ('点击 Renew 前', ARTIFACTS / '02-before-renew.png'),
    ('点击 Renew 后', ARTIFACTS / '03-after-renew-click.png'),
]


def main():
    doc = Document()
    doc.add_heading('HidenCloud Renew 执行报告', level=1)
    doc.add_paragraph(f'生成时间：{datetime.now().strftime("%F %T")}')
    doc.add_paragraph('目标：在指定日期自动打开目标服务页，执行 Renew，并留存前后证据图，用于后续形成每周执行计划。')

    doc.add_heading('关键信息', level=2)
    doc.add_paragraph('目标页面：https://dash.hidencloud.com/service/207229/manage')
    doc.add_paragraph('预期 Due date：20 Apr 2026')
    doc.add_paragraph('执行策略：先截图守门，确认目标页与 Due date，再点击 Renew，最后截图留证。')

    doc.add_heading('执行截图', level=2)
    for title, path in IMAGES:
        doc.add_paragraph(title)
        if path.exists():
            doc.add_picture(str(path), width=Inches(6.5))
        else:
            doc.add_paragraph(f'[缺失图片] {path.name}')

    doc.add_heading('执行日志', level=2)
    if LOG_PATH.exists():
        text = LOG_PATH.read_text(encoding='utf-8', errors='ignore')[-12000:]
        doc.add_paragraph(text)
    else:
        doc.add_paragraph('[缺失日志] run.log')

    doc.add_heading('后续建议', level=2)
    doc.add_paragraph('1. 观察点击 Renew 后页面真实变化，确认成功后是否出现新的按钮/弹窗/跳转。')
    doc.add_paragraph('2. 若本次报告确认可行，下次可据此固化每周一次的自动执行计划。')
    doc.add_paragraph('3. 若成功后页面样式稳定，可进一步把成功判定写死，减少人工复核。')

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_PATH))
    print(DOCX_PATH)


if __name__ == '__main__':
    main()
